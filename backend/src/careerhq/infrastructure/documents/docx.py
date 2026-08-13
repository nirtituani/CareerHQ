"""DOCX text extraction, via python-docx (MIT)."""

from __future__ import annotations

import io

from docx import Document


def extract(data: bytes) -> str:
    """Return the document's text, including table cells.

    Tables matter more than they might seem: a great many CV templates lay out
    contact details, skills or dates in a borderless table, and reading only
    paragraphs silently loses them — producing an extraction that looks
    plausible while missing exactly the fields a reviewer would assume were
    simply absent from the CV.
    """
    document = Document(io.BytesIO(data))

    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)

    return "\n".join(part for part in parts if part.strip()).strip()
